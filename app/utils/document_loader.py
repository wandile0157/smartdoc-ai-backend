"""
Document Loader - Extract text from various document formats
Supports: .txt, .pdf, .docx, .doc
"""

from typing import Optional
import io
import logging
from pathlib import Path

# Document processing imports
from docx import Document
import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Utility class for loading and extracting text from various document formats.
    """
    
    SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.docx', '.doc']
    MAX_FILE_SIZE_MB = 10
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            filename (str): Name of the file
            
        Returns:
            bool: True if supported
        """
        ext = Path(filename).suffix.lower()
        return ext in DocumentLoader.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def validate_file_size(file_size: int) -> bool:
        """
        Validate file size is within limits.
        
        Args:
            file_size (int): Size in bytes
            
        Returns:
            bool: True if valid
        """
        max_bytes = DocumentLoader.MAX_FILE_SIZE_MB * 1024 * 1024
        return file_size <= max_bytes
    
    @staticmethod
    def load_from_file(file_path: str) -> str:
        """
        Load text from a file path.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not DocumentLoader.is_supported(path.name):
            raise ValueError(f"Unsupported file type: {path.suffix}")
        
        # Check file size
        file_size = path.stat().st_size
        if not DocumentLoader.validate_file_size(file_size):
            raise ValueError(f"File too large. Maximum size: {DocumentLoader.MAX_FILE_SIZE_MB}MB")
        
        # Load based on extension
        ext = path.suffix.lower()
        
        try:
            if ext == '.txt':
                return DocumentLoader._load_txt(path)
            elif ext == '.pdf':
                return DocumentLoader._load_pdf(path)
            elif ext in ['.docx', '.doc']:
                return DocumentLoader._load_docx(path)
            else:
                raise ValueError(f"Unsupported extension: {ext}")
        
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {str(e)}")
            raise ValueError(f"Failed to load document: {str(e)}")
    
    @staticmethod
    def load_from_bytes(file_bytes: bytes, filename: str) -> str:
        """
        Load text from file bytes (for uploaded files).
        
        Args:
            file_bytes (bytes): File content as bytes
            filename (str): Original filename
            
        Returns:
            str: Extracted text
        """
        if not DocumentLoader.is_supported(filename):
            raise ValueError(f"Unsupported file type: {Path(filename).suffix}")
        
        # Check file size
        if not DocumentLoader.validate_file_size(len(file_bytes)):
            raise ValueError(f"File too large. Maximum size: {DocumentLoader.MAX_FILE_SIZE_MB}MB")
        
        # Load based on extension
        ext = Path(filename).suffix.lower()
        
        try:
            if ext == '.txt':
                return DocumentLoader._load_txt_from_bytes(file_bytes)
            elif ext == '.pdf':
                return DocumentLoader._load_pdf_from_bytes(file_bytes)
            elif ext in ['.docx', '.doc']:
                return DocumentLoader._load_docx_from_bytes(file_bytes)
            else:
                raise ValueError(f"Unsupported extension: {ext}")
        
        except Exception as e:
            logger.error(f"Error loading file {filename}: {str(e)}")
            raise ValueError(f"Failed to load document: {str(e)}")
    
    # ==================== TEXT FILE LOADERS ====================
    
    @staticmethod
    def _load_txt(file_path: Path) -> str:
        """Load text from .txt file"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    @staticmethod
    def _load_txt_from_bytes(file_bytes: bytes) -> str:
        """Load text from .txt bytes"""
        try:
            # Try UTF-8 first
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            return file_bytes.decode('latin-1')
    
    # ==================== PDF LOADERS ====================
    
    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        """Load text from PDF file using pdfplumber (better extraction)"""
        text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            if not text_parts:
                # Fallback to PyPDF2 if pdfplumber fails
                return DocumentLoader._load_pdf_pypdf2(file_path)
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            return DocumentLoader._load_pdf_pypdf2(file_path)
    
    @staticmethod
    def _load_pdf_from_bytes(file_bytes: bytes) -> str:
        """Load text from PDF bytes using pdfplumber"""
        text_parts = []
        
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            if not text_parts:
                # Fallback to PyPDF2
                return DocumentLoader._load_pdf_pypdf2_from_bytes(file_bytes)
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            return DocumentLoader._load_pdf_pypdf2_from_bytes(file_bytes)
    
    @staticmethod
    def _load_pdf_pypdf2(file_path: Path) -> str:
        """Fallback PDF loader using PyPDF2"""
        text_parts = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        if not text_parts:
            raise ValueError("Could not extract text from PDF")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _load_pdf_pypdf2_from_bytes(file_bytes: bytes) -> str:
        """Fallback PDF loader using PyPDF2 from bytes"""
        text_parts = []
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        if not text_parts:
            raise ValueError("Could not extract text from PDF")
        
        return '\n\n'.join(text_parts)
    
    # ==================== DOCX LOADERS ====================
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Load text from .docx file"""
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        if not text_parts:
            raise ValueError("No text content found in document")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _load_docx_from_bytes(file_bytes: bytes) -> str:
        """Load text from .docx bytes"""
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        if not text_parts:
            raise ValueError("No text content found in document")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def get_file_info(filename: str, file_size: int) -> dict:
        """
        Get information about a file.
        
        Args:
            filename (str): Name of the file
            file_size (int): Size in bytes
            
        Returns:
            dict: File information
        """
        ext = Path(filename).suffix.lower()
        
        return {
            'filename': filename,
            'extension': ext,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'is_supported': ext in DocumentLoader.SUPPORTED_EXTENSIONS,
            'is_valid_size': file_size <= (DocumentLoader.MAX_FILE_SIZE_MB * 1024 * 1024)
        }